# -*- coding: utf-8 -*-
"""
生成复杂的测试文件：政策文件.docx 和 配置表.xlsx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
import os

output_dir = r"c:\Users\LAKEWOOD\Desktop\毕业设计\marketing_audit_system\data\input"

def create_policy_document():
    doc = Document()
    
    title = doc.add_heading('企业营销活动管理规范', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('版本号：V3.0')
    doc.add_paragraph('生效日期：2024年1月1日')
    doc.add_paragraph('发布部门：市场营销管理部')
    doc.add_paragraph()
    
    doc.add_heading('第一章 总则', level=1)
    
    doc.add_heading('第一条 目的', level=2)
    doc.add_paragraph('为规范公司营销活动管理，防范营销风险，保护消费者合法权益，提升品牌形象，根据《中华人民共和国广告法》、《消费者权益保护法》及相关法律法规，结合公司实际情况，特制定本规范。')
    
    doc.add_heading('第二条 适用范围', level=2)
    doc.add_paragraph('本规范适用于公司及下属各分子公司、办事处开展的所有营销活动，包括但不限于：')
    items = [
        '优惠券、代金券、折扣券等各类券类促销活动',
        '满减、满赠、买赠等促销活动',
        '会员积分、会员专享等会员营销活动',
        '节日促销、周年庆等主题营销活动',
        '线上渠道（APP、小程序、官网）营销活动',
        '线下门店、专柜等实体渠道营销活动',
        '跨界合作、联合营销等合作推广活动'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('第三条 基本原则', level=2)
    doc.add_paragraph('营销活动应当遵循以下基本原则：')
    doc.add_paragraph('（一）合法合规原则：严格遵守国家法律法规，不得从事违法违规营销活动。')
    doc.add_paragraph('（二）诚信经营原则：营销活动内容应当真实、准确，不得虚假宣传、误导消费者。')
    doc.add_paragraph('（三）公平公正原则：营销活动规则应当公平合理，不得设置不公平、不合理的交易条件。')
    doc.add_paragraph('（四）风险可控原则：营销活动应当进行风险评估，确保活动风险可控。')
    
    doc.add_heading('第二章 优惠券管理', level=1)
    
    doc.add_heading('第四条 优惠券类型', level=2)
    doc.add_paragraph('公司优惠券分为以下类型：')
    
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['优惠券类型', '面额范围', '适用场景', '审批级别']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data = [
        ['新人专享券', '10-100元', '新用户注册', '运营经理'],
        ['满减优惠券', '20-500元', '日常促销', '运营经理'],
        ['折扣优惠券', '5%-50%', '会员专享', '运营总监'],
        ['免单优惠券', '最高1000元', '大型活动', 'VP及以上'],
        ['组合优惠券', '按需设定', '特殊活动', 'VP及以上']
    ]
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('第五条 优惠券金额限制', level=2)
    doc.add_paragraph('优惠券金额应当符合以下限制：')
    doc.add_paragraph('（一）单张优惠券面额不得超过500元，特殊促销活动经VP审批后可提高至1000元。')
    doc.add_paragraph('（二）同一用户单次活动领取优惠券总面额不得超过2000元。')
    doc.add_paragraph('（三）月度优惠券发放总金额不得超过当月营销预算的30%。')
    doc.add_paragraph('（四）年度优惠券发放总金额不得超过年度营销预算的25%。')
    doc.add_paragraph('（五）优惠券折扣比例不得低于商品成本的80%，防止恶意套利。')
    
    doc.add_heading('第六条 优惠券发放规则', level=2)
    doc.add_paragraph('优惠券发放应当遵循以下规则：')
    doc.add_paragraph('（一）新人专享券仅限注册时间在30天内的新用户领取，每人限领1次。')
    doc.add_paragraph('（二）普通优惠券每位用户每月领取次数不得超过3次。')
    doc.add_paragraph('（三）VIP用户每月领取次数可放宽至5次，需验证VIP身份。')
    doc.add_paragraph('（四）优惠券发放对象必须明确限定，不得向不符合条件的用户发放。')
    doc.add_paragraph('（五）禁止通过任何形式买卖、转让优惠券，一经发现立即作废。')
    
    doc.add_heading('第七条 优惠券有效期', level=2)
    doc.add_paragraph('优惠券有效期管理要求：')
    doc.add_paragraph('（一）优惠券有效期不得少于7天，给予用户充足使用时间。')
    doc.add_paragraph('（二）优惠券有效期不得超过90天，防止长期占用营销资源。')
    doc.add_paragraph('（三）新人专享券有效期应当为7-30天。')
    doc.add_paragraph('（四）活动专属券有效期应当与活动周期一致，活动结束后自动失效。')
    doc.add_paragraph('（五）优惠券到期前3天应当通过短信或APP推送提醒用户。')
    
    doc.add_heading('第三章 促销活动管理', level=1)
    
    doc.add_heading('第八条 活动预算管理', level=2)
    doc.add_paragraph('促销活动预算应当符合以下规定：')
    doc.add_paragraph('（一）单次促销活动总预算不得超过100万元。')
    doc.add_paragraph('（二）季度促销活动累计预算不得超过300万元。')
    doc.add_paragraph('（三）年度促销活动累计预算不得超过1000万元。')
    doc.add_paragraph('（四）预算超过50万元的活动需提前15个工作日提交审批。')
    doc.add_paragraph('（五）预算超过100万元的活动需提前30个工作日提交审批，并附详细ROI分析报告。')
    
    doc.add_heading('第九条 活动范围限制', level=2)
    doc.add_paragraph('促销活动范围应当符合以下限制：')
    doc.add_paragraph('（一）线上专属活动仅限线上渠道（APP、小程序、官网）参与，禁止线下门店参与。')
    doc.add_paragraph('（二）线下门店活动需经区域经理审批后方可开展。')
    doc.add_paragraph('（三）全国性活动需经总部营销中心审批。')
    doc.add_paragraph('（四）区域性活动范围不得超过3个省级行政区。')
    doc.add_paragraph('（五）跨界合作活动需经法务部门审核合作协议后方可开展。')
    
    doc.add_heading('第十条 活动叠加规则', level=2)
    doc.add_paragraph('促销活动叠加应当遵循以下规则：')
    doc.add_paragraph('（一）同一订单不得同时使用两种及以上优惠券。')
    doc.add_paragraph('（二）满减活动与折扣活动不得叠加使用。')
    doc.add_paragraph('（三）会员折扣与活动折扣不得叠加，取优惠力度较大者。')
    doc.add_paragraph('（四）积分抵现与优惠券可叠加使用，但积分抵现比例不得超过订单金额的20%。')
    doc.add_paragraph('（五）禁止设置隐性消费门槛，所有使用条件必须在活动页面显著位置明示。')
    
    doc.add_heading('第四章 用户权益保护', level=1)
    
    doc.add_heading('第十一条 信息披露要求', level=2)
    doc.add_paragraph('营销活动信息披露应当符合以下要求：')
    doc.add_paragraph('（一）活动规则必须在活动开始前24小时向用户公示。')
    doc.add_paragraph('（二）活动规则应当包含：活动时间、参与条件、优惠内容、使用规则、有效期等。')
    doc.add_paragraph('（三）限制性条款应当以醒目方式标注，不得使用模糊表述。')
    doc.add_paragraph('（四）活动规则变更需提前7天通知用户，已发放优惠券按原规则执行。')
    doc.add_paragraph('（五）活动结束后7天内应当公布活动中奖名单（如涉及）。')
    
    doc.add_heading('第十二条 消费门槛限制', level=2)
    doc.add_paragraph('消费门槛设置应当符合以下限制：')
    doc.add_paragraph('（一）优惠券使用门槛不得超过优惠券面额的5倍。')
    doc.add_paragraph('（二）满减活动门槛不得超过优惠金额的10倍。')
    doc.add_paragraph('（三）不得设置仅限特定商品使用的隐性门槛。')
    doc.add_paragraph('（四）门槛金额应当为整数，不得设置小数点后两位。')
    doc.add_paragraph('（五）门槛计算应当清晰透明，不得在结算环节临时增加费用。')
    
    doc.add_heading('第十三条 退款处理规则', level=2)
    doc.add_paragraph('涉及优惠活动的退款处理：')
    doc.add_paragraph('（一）使用优惠券的订单退款时，优惠券应当返还至用户账户。')
    doc.add_paragraph('（二）返还的优惠券有效期应当重新计算，不得短于原有效期。')
    doc.add_paragraph('（三）部分退款时，优惠券不返还，但可按比例退还差价。')
    doc.add_paragraph('（四）活动期间产生的积分在退款后应当扣除。')
    doc.add_paragraph('（五）因商品质量问题导致的退款，用户可选择返还优惠券或获得等值补偿。')
    
    doc.add_heading('第五章 风险控制', level=1)
    
    doc.add_heading('第十四条 风险评估要求', level=2)
    doc.add_paragraph('营销活动开展前应当进行风险评估：')
    doc.add_paragraph('（一）预算超过10万元的活动必须进行风险评估。')
    doc.add_paragraph('（二）风险评估应当包含：财务风险、法律风险、舆情风险、运营风险。')
    doc.add_paragraph('（三）高风险活动需制定应急预案。')
    doc.add_paragraph('（四）活动期间应当进行实时监控，发现异常及时处置。')
    doc.add_paragraph('（五）活动结束后应当进行复盘分析，总结经验教训。')
    
    doc.add_heading('第十五条 禁止行为', level=2)
    doc.add_paragraph('营销活动中禁止以下行为：')
    doc.add_paragraph('（一）虚假宣传、夸大宣传、误导消费者。')
    doc.add_paragraph('（二）先涨价后降价、虚构原价。')
    doc.add_paragraph('（三）设置不公平、不合理的交易条件。')
    doc.add_paragraph('（四）诱导用户进行非真实消费。')
    doc.add_paragraph('（五）泄露用户个人信息。')
    doc.add_paragraph('（六）恶意刷单、虚假交易。')
    doc.add_paragraph('（七）违反国家法律法规的其他行为。')
    
    doc.add_heading('第六章 审批流程', level=1)
    
    doc.add_heading('第十六条 审批权限', level=2)
    
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    headers2 = ['活动类型', '预算范围', '审批人', '审批时限']
    for i, header in enumerate(headers2):
        table2.rows[0].cells[i].text = header
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data2 = [
        ['日常促销', '≤10万元', '运营经理', '3个工作日'],
        ['中型活动', '10-50万元', '运营总监', '5个工作日'],
        ['大型活动', '50-100万元', 'VP', '7个工作日'],
        ['特批活动', '>100万元', 'CEO', '10个工作日']
    ]
    for i, row_data in enumerate(data2, 1):
        for j, cell_data in enumerate(row_data):
            table2.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('第七章 附则', level=1)
    
    doc.add_heading('第十七条 解释权', level=2)
    doc.add_paragraph('本规范由市场营销管理部负责解释。')
    
    doc.add_heading('第十八条 生效日期', level=2)
    doc.add_paragraph('本规范自发布之日起施行，原有相关规定同时废止。')
    
    doc.add_heading('第十九条 违规处理', level=2)
    doc.add_paragraph('违反本规范规定的，视情节轻重给予以下处理：')
    doc.add_paragraph('（一）情节轻微的，给予口头警告，责令限期整改。')
    doc.add_paragraph('（二）情节较重的，给予书面警告，扣减绩效奖金。')
    doc.add_paragraph('（三）情节严重的，给予降职降薪处理。')
    doc.add_paragraph('（四）造成重大损失的，解除劳动合同，追究法律责任。')
    
    filepath = os.path.join(output_dir, '营销活动管理规范.docx')
    doc.save(filepath)
    print(f"政策文件已生成: {filepath}")
    return filepath


def create_config_excel():
    wb = openpyxl.Workbook()
    
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')
    
    ws1 = wb.active
    ws1.title = '优惠券配置'
    
    headers1 = ['配置ID', '活动名称', '优惠券类型', '面额(元)', '使用门槛(元)', 
                '发放对象', '每人限领', '有效期(天)', '发放总量', '已发放量',
                '活动渠道', '是否可叠加', '审批状态', '创建时间']
    
    for col, header in enumerate(headers1, 1):
        cell = ws1.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    
    coupon_data = [
        ['CPN_2024_001', '春节大促-新人券', '新人专享券', 150, 500, '全部用户', 5, 90, 100000, 85000, '全渠道', '是', '已生效', '2024-01-15'],
        ['CPN_2024_002', '春节大促-满减券', '满减优惠券', 600, 2000, '全部用户', 3, 60, 50000, 42000, '全渠道', '是', '已生效', '2024-01-15'],
        ['CPN_2024_003', 'VIP专属折扣', '折扣优惠券', 50, 0, 'VIP用户', 10, 30, 20000, 18000, 'APP/小程序', '否', '已生效', '2024-01-20'],
        ['CPN_2024_004', '周年庆免单券', '免单优惠券', 1200, 0, '全部用户', 1, 7, 1000, 1000, '全渠道', '否', '已生效', '2024-02-01'],
        ['CPN_2024_005', '会员日专享', '满减优惠券', 300, 800, '会员用户', 2, 14, 30000, 25000, 'APP', '是', '已生效', '2024-02-10'],
        ['CPN_2024_006', '新用户礼包', '组合优惠券', 200, 300, '新注册用户', 1, 30, 50000, 38000, '全渠道', '是', '已生效', '2024-02-15'],
        ['CPN_2024_007', '清仓特惠券', '折扣优惠券', 70, 0, '全部用户', 5, 15, 80000, 72000, '线下门店', '是', '已生效', '2024-02-20'],
        ['CPN_2024_008', '跨品类组合券', '组合优惠券', 400, 1000, '全部用户', 3, 45, 15000, 12000, 'APP/小程序', '否', '待审批', '2024-02-25'],
    ]
    
    for row_idx, row_data in enumerate(coupon_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws1.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')
    
    for col in range(1, len(headers1) + 1):
        ws1.column_dimensions[get_column_letter(col)].width = 15
    
    ws2 = wb.create_sheet('活动预算配置')
    
    headers2 = ['活动ID', '活动名称', '活动类型', '总预算(元)', '已使用预算(元)',
                '预算审批人', '审批级别', '活动开始日期', '活动结束日期', 
                '活动范围', '目标用户群', '预期ROI', '实际ROI', '风险等级']
    
    for col, header in enumerate(headers2, 1):
        cell = ws2.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    
    budget_data = [
        ['ACT_2024_001', '春节大促活动', '大型促销', 2000000, 1500000, '张总监', 'VP', '2024-01-20', '2024-02-20', '全国', '全部用户', '3.5', '2.8', '高'],
        ['ACT_2024_002', 'VIP会员日', '会员营销', 300000, 280000, '李经理', '运营经理', '2024-02-01', '2024-02-28', '线上', 'VIP用户', '4.0', '4.2', '低'],
        ['ACT_2024_003', '周年庆典', '主题营销', 1500000, 800000, '王VP', 'CEO', '2024-03-01', '2024-03-15', '全国', '全部用户', '5.0', '待统计', '中'],
        ['ACT_2024_004', '新用户拉新', '获客活动', 500000, 450000, '赵经理', '运营总监', '2024-01-01', '2024-03-31', '全渠道', '新用户', '2.0', '1.8', '中'],
        ['ACT_2024_005', '清仓特卖', '库存清理', 200000, 180000, '孙经理', '运营经理', '2024-02-15', '2024-02-28', '线下门店', '全部用户', '1.5', '1.6', '低'],
        ['ACT_2024_006', '跨品牌联合促销', '合作营销', 800000, 600000, '周总监', 'VP', '2024-03-10', '2024-03-25', '华东区', '全部用户', '3.0', '待统计', '中'],
    ]
    
    for row_idx, row_data in enumerate(budget_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws2.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')
    
    for col in range(1, len(headers2) + 1):
        ws2.column_dimensions[get_column_letter(col)].width = 15
    
    ws3 = wb.create_sheet('活动规则配置')
    
    headers3 = ['规则ID', '关联活动', '规则类型', '规则内容', '约束条件',
                '适用商品', '排除商品', '生效时间', '失效时间', '规则状态']
    
    for col, header in enumerate(headers3, 1):
        cell = ws3.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    
    rule_data = [
        ['RULE_001', 'ACT_2024_001', '门槛限制', '满500减100', '订单金额≥500元', '全部商品', '虚拟商品', '2024-01-20', '2024-02-20', '生效中'],
        ['RULE_002', 'ACT_2024_001', '叠加规则', '可与积分叠加', '积分抵现≤20%', '全部商品', '无', '2024-01-20', '2024-02-20', '生效中'],
        ['RULE_003', 'ACT_2024_001', '范围限制', '仅限线上渠道', 'APP/小程序/官网', '全部商品', '无', '2024-01-20', '2024-02-20', '生效中'],
        ['RULE_004', 'ACT_2024_002', '身份限制', '仅限VIP用户', 'VIP等级≥3级', '全部商品', '无', '2024-02-01', '2024-02-28', '生效中'],
        ['RULE_005', 'ACT_2024_002', '折扣上限', '最高折扣500元', '单笔订单', '全部商品', '无', '2024-02-01', '2024-02-28', '生效中'],
        ['RULE_006', 'ACT_2024_003', '时间限制', '每日10:00-22:00', '限时抢购', '指定商品', '无', '2024-03-01', '2024-03-15', '待生效'],
        ['RULE_007', 'ACT_2024_003', '数量限制', '每人限购3件', '单品限量', '活动商品', '无', '2024-03-01', '2024-03-15', '待生效'],
        ['RULE_008', 'ACT_2024_004', '身份限制', '仅限新注册用户', '注册时间≤30天', '全部商品', '无', '2024-01-01', '2024-03-31', '生效中'],
        ['RULE_009', 'ACT_2024_004', '首单限制', '仅限首单使用', '无历史订单', '全部商品', '无', '2024-01-01', '2024-03-31', '生效中'],
        ['RULE_010', 'ACT_2024_005', '渠道限制', '仅限线下门店', '指定门店', '清仓商品', '新品', '2024-02-15', '2024-02-28', '生效中'],
    ]
    
    for row_idx, row_data in enumerate(rule_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws3.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')
    
    for col in range(1, len(headers3) + 1):
        ws3.column_dimensions[get_column_letter(col)].width = 18
    
    ws4 = wb.create_sheet('用户权益配置')
    
    headers4 = ['权益ID', '权益名称', '权益类型', '权益内容', '适用用户',
                '生效条件', '使用次数限制', '有效期类型', '权益状态', '备注']
    
    for col, header in enumerate(headers4, 1):
        cell = ws4.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    
    benefit_data = [
        ['BF_001', '新人礼包', '优惠券包', '含3张优惠券', '新注册用户', '完成注册', '1次', '固定期限', '生效中', '含满减券、折扣券、免邮券'],
        ['BF_002', '会员折扣', '价格优惠', '全场95折', '会员用户', '会员等级≥1', '不限', '长期有效', '生效中', '与活动折扣取最优'],
        ['BF_003', 'VIP专享价', '价格优惠', '指定商品9折', 'VIP用户', 'VIP等级≥3', '不限', '长期有效', '生效中', '部分商品除外'],
        ['BF_004', '积分加倍', '积分权益', '购物积分2倍', 'VIP用户', 'VIP等级≥5', '不限', '活动期间', '生效中', '仅限活动商品'],
        ['BF_005', '免费退换', '服务权益', '7天无理由退换', '全部用户', '无', '不限', '长期有效', '生效中', '商品需保持原样'],
        ['BF_006', '专属客服', '服务权益', '一对一客服服务', 'VIP用户', 'VIP等级≥3', '不限', '长期有效', '生效中', '工作日9:00-18:00'],
    ]
    
    for row_idx, row_data in enumerate(benefit_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws4.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')
    
    for col in range(1, len(headers4) + 1):
        ws4.column_dimensions[get_column_letter(col)].width = 18
    
    ws5 = wb.create_sheet('风险控制配置')
    
    headers5 = ['风控ID', '风控类型', '监控指标', '阈值设置', '触发条件',
                '处理方式', '通知对象', '监控频率', '风控状态', '最后检查时间']
    
    for col, header in enumerate(headers5, 1):
        cell = ws5.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    
    risk_data = [
        ['RC_001', '预算超支', '预算使用率', '≥90%', '实时监控', '自动预警', '活动负责人', '实时', '生效中', '2024-02-28 10:00'],
        ['RC_002', '异常领取', '单用户领取量', '≥10张/天', '实时监控', '限制领取', '风控专员', '实时', '生效中', '2024-02-28 10:00'],
        ['RC_003', '恶意刷单', '订单取消率', '≥30%', '每小时统计', '人工审核', '风控主管', '每小时', '生效中', '2024-02-28 09:00'],
        ['RC_004', '舆情风险', '负面评价数', '≥100条/天', '每日统计', '启动预案', '公关部门', '每日', '生效中', '2024-02-28 08:00'],
        ['RC_005', '系统异常', '接口响应时间', '≥3秒', '实时监控', '自动降级', '技术团队', '实时', '生效中', '2024-02-28 10:00'],
        ['RC_006', '库存风险', '库存周转率', '≤0.5', '每日统计', '预警通知', '采购部门', '每日', '生效中', '2024-02-28 08:00'],
    ]
    
    for row_idx, row_data in enumerate(risk_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws5.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')
    
    for col in range(1, len(headers5) + 1):
        ws5.column_dimensions[get_column_letter(col)].width = 15
    
    filepath = os.path.join(output_dir, '营销活动配置表.xlsx')
    wb.save(filepath)
    print(f"配置表已生成: {filepath}")
    return filepath


if __name__ == '__main__':
    print("开始生成测试文件...")
    print()
    
    policy_file = create_policy_document()
    config_file = create_config_excel()
    
    print()
    print("=" * 60)
    print("文件生成完成！")
    print("=" * 60)
    print(f"政策文件: {policy_file}")
    print(f"配置表: {config_file}")
