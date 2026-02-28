import re
from typing import Dict, List, Any, Optional
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


class DocxParser:
    def __init__(self):
        self.current_section = None
        self.sections = {}
    
    def parse(self, filepath: str) -> Dict[str, Any]:
        doc = Document(filepath)
        result = {
            "metadata": self._extract_metadata(doc),
            "content": self._extract_content(doc),
            "tables": self._extract_tables(doc),
            "structure": self._analyze_structure(doc),
            "raw_text": self._get_raw_text(doc)
        }
        return result
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        metadata = {
            "title": "",
            "author": "",
            "created_date": "",
            "modified_date": "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
        
        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created_date"] = str(core_props.created)
        if core_props.modified:
            metadata["modified_date"] = str(core_props.modified)
        
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            if text and len(text) < 100 and not metadata["title"]:
                if any(kw in text for kw in ["通知", "办法", "规定", "方案", "政策"]):
                    metadata["title"] = text
                    break
        
        return metadata
    
    def _extract_content(self, doc: Document) -> List[Dict[str, Any]]:
        content = []
        current_section = "正文"
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            para_info = {
                "index": i,
                "text": text,
                "style": para.style.name if para.style else None,
                "is_heading": self._is_heading(para),
                "section": current_section
            }
            
            if para_info["is_heading"]:
                current_section = text
                para_info["section"] = "标题"
            
            content.append(para_info)
        
        return content
    
    def _is_heading(self, para: Paragraph) -> bool:
        if not para.style:
            return False
        style_name = para.style.name.lower()
        if "heading" in style_name or "标题" in style_name:
            return True
        text = para.text.strip()
        heading_patterns = [
            r'^[一二三四五六七八九十]+[、．.]',
            r'^\d+[、．.]',
            r'^\([一二三四五六七八九十]+\)',
            r'^\(\d+\)',
            r'^第[一二三四五六七八九十]+[章节条款]',
        ]
        for pattern in heading_patterns:
            if re.match(pattern, text):
                return True
        return False
    
    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = self._parse_table(table, i)
            tables.append(table_data)
        return tables
    
    def _parse_table(self, table: Table, table_index: int) -> Dict[str, Any]:
        rows_data = []
        merged_cells = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append(cell_text)
                
                if cell._tc.get('vMerge') is not None:
                    merged_cells.append({
                        "row": row_idx,
                        "col": col_idx,
                        "type": "vertical_merge"
                    })
            
            rows_data.append(row_data)
        
        headers = rows_data[0] if rows_data else []
        
        return {
            "table_index": table_index,
            "headers": headers,
            "rows": rows_data,
            "row_count": len(rows_data),
            "col_count": len(headers) if headers else 0,
            "merged_cells": merged_cells,
            "markdown": self._table_to_markdown(rows_data)
        }
    
    def _table_to_markdown(self, rows: List[List[str]]) -> str:
        if not rows:
            return ""
        
        lines = []
        for i, row in enumerate(rows):
            line = "| " + " | ".join(str(cell).replace('\n', ' ') for cell in row) + " |"
            lines.append(line)
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(row)) + " |"
                lines.append(separator)
        
        return "\n".join(lines)
    
    def _analyze_structure(self, doc: Document) -> Dict[str, Any]:
        structure = {
            "sections": [],
            "outline": []
        }
        
        for para in doc.paragraphs:
            if self._is_heading(para):
                text = para.text.strip()
                level = self._get_heading_level(para)
                structure["sections"].append({
                    "title": text,
                    "level": level
                })
                structure["outline"].append(f"{'  ' * (level - 1)}{text}")
        
        return structure
    
    def _get_heading_level(self, para: Paragraph) -> int:
        if not para.style:
            return 1
        style_name = para.style.name
        if "Heading" in style_name:
            match = re.search(r'\d+', style_name)
            if match:
                return int(match.group())
        return 1
    
    def _get_raw_text(self, doc: Document) -> str:
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    
    def to_json_format(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "document_type": "policy_document",
            "metadata": parsed_data["metadata"],
            "structure": {
                "sections": parsed_data["structure"]["sections"],
                "outline": parsed_data["structure"]["outline"]
            },
            "content": [
                {
                    "section": item["section"],
                    "text": item["text"],
                    "type": "heading" if item["is_heading"] else "paragraph"
                }
                for item in parsed_data["content"]
            ],
            "tables": [
                {
                    "index": t["table_index"],
                    "headers": t["headers"],
                    "data": t["rows"],
                    "markdown": t["markdown"]
                }
                for t in parsed_data["tables"]
            ],
            "raw_text": parsed_data["raw_text"]
        }
